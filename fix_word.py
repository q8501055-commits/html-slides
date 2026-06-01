import re
import os
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc_path = r"C:\Users\User\.openclaw\workspace\Humanoid_Robot_Tactile_Sensors_2026_Final.docx"
content_path = r"C:\Users\User\.openclaw\workspace\content.md"

def read_markdown(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

def create_styled_word_document(markdown_text, output_path):
    doc = docx.Document()
    
    # Customizing default styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft JhengHei'
    font.size = Pt(12)
    
    # Check if Heading 1 exists, if not just use normal
    try:
        for i in range(1, 4):
            h_style = doc.styles[f'Heading {i}']
            h_font = h_style.font
            h_font.name = 'Microsoft JhengHei'
            h_font.color.rgb = RGBColor(0, 0, 0)
            if i == 1:
                h_font.size = Pt(18)
                h_font.bold = True
                h_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif i == 2:
                h_font.size = Pt(16)
                h_font.bold = True
            elif i == 3:
                h_font.size = Pt(14)
                h_font.bold = True
    except KeyError:
        pass
    
    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            doc.add_paragraph(line, style='List Number')
        else:
            doc.add_paragraph(line)
            
    doc.save(output_path)

markdown_text = read_markdown(content_path)
create_styled_word_document(markdown_text, doc_path)
print("Word document generated with proper headings.")